import re
from typing import Optional
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown
import logging

import logging
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class ContentParser:
    """Parse various content formats into plain text."""

    @staticmethod
    async def parse_content(file_path: str, content_type: str) -> str:
        """Parse content based on type."""
        try:
            if content_type == "html":
                return await ContentParser._parse_html(file_path)
            elif content_type == "markdown":
                return await ContentParser._parse_markdown(file_path)
            elif content_type == "pdf":
                return await ContentParser._parse_pdf(file_path)
            elif content_type == "docx":
                return await ContentParser._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

        except Exception as e:
            logger.error(f"Error parsing {content_type}: {str(e)}")
            raise

    @staticmethod
    async def _parse_html(file_path: str) -> str:
        """Parse HTML file to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)

        return text

    @staticmethod
    async def _parse_markdown(file_path: str) -> str:
        """Parse Markdown file to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert to HTML first, then to text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()

        return text.strip()

    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """Parse PDF file to plain text."""
        text = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())

        return '\n'.join(text)

    @staticmethod
    def _iter_block_items(parent):
        """
        Generate a reference to each paragraph and table child within parent, in document order.
        Each returned value is an instance of either Table or Paragraph.
        """
        if isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            parent_elm = parent.element.body

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """Parse DOCX file to plain text including tables."""
        doc = docx.Document(file_path)
        text = []
        
        for block in ContentParser._iter_block_items(doc):
            if isinstance(block, Paragraph):
                text.append(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    row_text = []
                    for cell in row.cells:
                        # Recursively extract text from cell if needed, but for now simple text access
                        # Cells can contain paragraphs too, but cell.text gets all of it joined
                        row_text.append(cell.text.strip())
                    text.append(" | ".join(row_text))
        
        return '\n'.join(text)


class ContentParserService:
    """
    Synchronous wrapper for ContentParser.
    Used by PreprocessingService which needs sync methods.
    """
    
    def parse_pdf(self, file_path: str) -> str:
        """Synchronously parse PDF file."""
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def extract_pdf_pages(self, file_path: str) -> list[str]:
        """
        Extract PDF as list of page texts for page-aware chunking.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of page texts (one string per page)
        """
        pages = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                pages.append(page_text if page_text else "")
        return pages
    
    def parse_docx(self, file_path: str) -> str:
        """Synchronously parse DOCX file."""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)


content_parser = ContentParser()
