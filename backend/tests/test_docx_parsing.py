
import pytest
import os
import docx
from app.services.content_parser import ContentParser

@pytest.mark.asyncio
async def test_parse_docx_with_tables():
    """Test parsing a DOCX file containing both paragraphs and tables."""
    filename = "test_table_parsing.docx"
    
    # Create test file
    doc = docx.Document()
    doc.add_paragraph("Paragraph 1 before table.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell 1"
    table.cell(1, 1).text = "Cell 2"
    
    doc.add_paragraph("Paragraph 2 after table.")
    doc.save(filename)
    
    try:
        # Parse content
        content = await ContentParser._parse_docx(filename)
        
        print(f"\nParsed Content:\n{content}")
        
        # Verify content exists and is in order
        assert "Paragraph 1 before table." in content
        assert "Header 1" in content
        assert "Header 2" in content
        assert "Cell 1" in content
        assert "Cell 2" in content
        assert "Paragraph 2 after table." in content
        
        # Check order (heuristic)
        p1_idx = content.find("Paragraph 1 before table.")
        h1_idx = content.find("Header 1")
        p2_idx = content.find("Paragraph 2 after table.")
        
        assert p1_idx < h1_idx, "Table content should come after first paragraph"
        assert h1_idx < p2_idx, "Table content should come before second paragraph"
        
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    import asyncio
    asyncio.run(test_parse_docx_with_tables())
